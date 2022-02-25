from distutils.command.upload import upload
from multiprocessing import context
from django.shortcuts import render
from .models import UploadPDF
from .forms import UploadPDFForm
from PIL import Image
import pytesseract
import sys
from pdf2image import convert_from_path
import os
from os import listdir
from os.path import isfile, join
import pandas as pd
import csv
import aspose.words as aw
from docx import Document
from docx.table import _Cell
from django.core.files.storage import default_storage
import re
import glob
from datetime import datetime

# Create your views here.
def Converter(request):

    context = {}

    if request.method == 'POST':
        print(request.FILES["upload"])
        print(request.FILES["upload"].content_type)
        form = {}
        # form = UploadPDFForm(request.POST, request.FILES)
        files = request.FILES.getlist('upload')
        allFiles = []
        for x in files:
            print(x)
            file_name = default_storage.save(x.name, x)
            allFiles.append(x.name)
        # if form.is_valid():
        #     files = form.cleaned_data["upload"]
        #     print(files)
        pytesseract.pytesseract.tesseract_cmd = r"C:/Program Files/Tesseract-OCR/tesseract.exe"
        # allFiles = [f for f in listdir("pdf/") if isfile(join("pdf/", f))]
        date_name = datetime.now().strftime("%Y%m%d%I%M%S%p")
        out_file = []
        for index, PDF_file in enumerate(allFiles):
            print("Converting Pdf File To Text: " + PDF_file[:-4])
            pdf_filepath = "documents/" + PDF_file
            pages = convert_from_path(
                pdf_filepath,
                200,
                grayscale=True,
                poppler_path=r"Release-22.01.0-0\poppler-22.01.0\Library\bin",
            )

            image_counter = 1
            for page in pages:
                filename = "page_" + str(image_counter) + ".jpg"
                page.save(filename, "JPEG")
                image_counter = image_counter + 1

            filelimit = image_counter - 1
            outfile = "out_text/" + PDF_file[:-4] + "_" + str(date_name) + ".txt"
            out_file.append(PDF_file[:-4] + "_" + str(date_name) + ".txt")
            
            if not os.path.exists("out_text/"):
                os.makedirs("out_text/")
            f = open(outfile, "w")
            for i in range(1, filelimit + 1):
                filename = "page_" + str(i) + ".jpg"
                text = str(((pytesseract.image_to_string(Image.open(filename)))))
                text = text.replace("-\n", "")
                f.write(text)
            f.close()
            print("Converted file: " + PDF_file[:-4])
        print(out_file)
        LicenceID = []
        Facility_Name = []
        country = []
        Telephone_Number = []
        date = []
        address = []
        name = []
        title = []
        Firm_Name = []
        address1 = []
        Telephone = []
        Percent = []
        Unrestricted = []
        INCOME_STATEMENT = []
        cash = []
        cash_ac = []
        cash_dep = []
        cash_dep_ac = []
        acc = []
        acc_ac = []
        supply_op = []
        supply_ac = []
        short_op = []
        short_ac = []
        prepaid_op = []
        prepaid_ac = []
        other_op = []
        other_ac = []
        account_op = []
        account_ac = []
        other_specify_op = []
        other_specify_ac = []
        current_assets_op = []
        current_assets_ac = []
        long_op = []
        long_ac = []
        long_term_op = []
        long_term_ac = []
        land_op = []
        land_ac = []
        build_op = []
        build_ac = []
        leasehold_op = []
        leasehold_ac = []
        equipment_op = []
        equipment_ac = []
        accumulated_op = []
        accumulated_ac = []
        deferred_op = []
        deferred_ac = []
        organization_op = []
        organization_ac = []
        accumulated_amortization_op = []
        accumulated_amortization_ac = []
        restricted_op = []
        restricted_ac = []
        other_long_term_op = []
        other_long_term_ac = []
        other_specify_two_op = []
        other_specify_two_ac = []
        long_term_assets_op = []
        long_term_assets_ac = []
        total_assets_op = []
        total_assets_ac = []
        accounts_payable_op = []
        accounts_payable_ac = []
        offer_op = []
        offer_ac = []
        accounts_patient_op = []
        accounts_patient_ac = []
        sort_term_op = []
        sort_term_ac = []
        accured_salaries_op = []
        accured_salaries_ac = []
        accrued_taxes_op = []
        accrued_taxes_ac = []
        accrued_real_op = []
        accrued_real_ac = []
        accrued_interest_op = []
        accrued_interest_ac = []
        deferred_compensation_op = []
        deferred_compensation_ac = []
        federal_op = []
        federal_ac = []
        current_liabilities_op = []
        current_liabilities_ac = []
        total_current_liabilities_op = []
        total_current_liabilities_ac = []
        long_term_notes_op = []
        long_term_notes_ac = []
        mortgage_op = []
        mortgage_ac = []
        bond_op = []
        bond_ac = []
        deferred_compensation_two_op = []
        deferred_compensation_two_ac = []
        long_liabilities_op = []
        long_liabilities_ac = []
        total_liabilities_op = []
        total_liabilities_ac = []
        total_equity_op = []
        total_equity_ac = []
        liabilities_equity_op = []
        liabilities_equity_ac = []
        gross = []
        disc = []
        sub_total = []
        day_care = []
        other_care = []
        therapy = []
        oxygen = []
        subtotal_ancillary = []
        payment = []
        other_gov = []
        can = []
        gift = []
        barber = []
        non_pat = []
        tele = []
        rental = []
        sale_drug = []
        sale_sup = []
        lab = []
        radio = []
        other_medi = []
        laundry = []
        subtotal_other = []
        # non_op = []
        contri = []
        interest = []
        subtotal_non_op = []
        settlement = []
        gov = []
        subtotal_other_revenue = []
        total_revenue = []
        general = []
        health = []
        general_admin = []
        ownership = []
        special_cost = []
        provider = []
        total_expenses = []
        income_before = []
        income_tax = []
        net_income = []
        madicaid = []
        private = []
        medicare = []
        other_specify_med = []
        other_speci = []
        total_inpa = []
        header = []
        listrow = []
        fileData = ""

        def check_list(base_list, counter):
            if len(base_list) < counter:
                base_list.append(0)
            return base_list

        file_counter = 1
        for i in out_file:
            with open("out_text/" + str(i)[:-4] + ".txt") as fp:
                fileData = fp.read()
                fp.close()

            # IDPH License ID Number
            a = (
                fileData.split("IDPH License ID Number: ")[1]
                .split(" ")[0]
                .replace("\n", "")
                .replace("Facility", "")
            )
            LicenceID.append(a)

            # Telephone Number
            d = fileData.split("Telephone Number: ")[1].split("F")[0].strip()
            Telephone_Number.append(d)

            # Address
            e = (
                fileData.split("Address: ")[1]
                .split(",")[0]
                .strip()
                .replace("State of Illinois", "")
            )
            address.append(e)

            # Telephone
            m = fileData.split("(Telephone) ")[1].split("F")[0].strip("_")
            Telephone.append(m)

            f = (
                fileData.split("Date of Initial License for Current Owners:")[1]
                .strip()
                .split(" ")[0]
                .replace("\n", "")
                .replace("Officer", "")
            )
            date.append(f)

            file1 = open("out_text/" + str(i)[:-4] + ".txt")
            Lines = file1.readlines()
            file1.close()

            
            other_specific_counter = 1
            for line in Lines:
                
                # Facility Name
                if "Facility Name:" in line:
                    b = line.split("Facility Name:")[1].strip()
                    # print(line.split('Facility Name:')[1])
                    Facility_Name.append(b)

                # County
                if "County:" in line:
                    c = line.split("County: ")[1].split(" ")[0].strip()
                    country.append(c)

                # Print Name
                if "(Print Name" in line:
                    g = line.split("(Print Name ")
                    if len(g) > 1:
                        g = g[1].strip().replace("—","").replace("_","")
                        #print(g)
                        name.append(g)

                # # Title
                if "and Title)" in line:
                    h = line.split("and Title) ")
                    if len(h) > 1:
                        h = h[-1].split(" ")[0].strip().replace("|__|","")
                        title.append(h)

                # # Firm Name
                if "(Firm Name" in line:
                    j = line.split("(Firm Name ")
                    if len(j) > 1:
                        j = j[1].strip().replace("—","").replace("_","")
                        Firm_Name.append(j)

                # # Address
                if "& Address)" in line:
                    k = line.split("& Address) ")
                    if len(k) > 1:
                        k = k[1].strip()
                        address1.append(k)

                if "column 4.)" in line:
                    n = (
                        fileData.split("column 4.) ")[-1]
                        .split(" ")[0]
                        .strip()
                        .replace("\n", "")
                        .replace("HFS", "")
                        .replace("2020STATE", "")
                    )
                    Percent.append(n)
                    # print(Percent)
                # XV. BALANCE SHEET
                if "XV. BALANCE SHEET" in line:
                    o = "XV. BALANCE SHEET"
                    Unrestricted.append(o)

                # XVII. INCOME STATEMENT
                if "XVII. INCOME STATEMENT" in line:
                    p = "XVII. INCOME STATEMENT"
                    INCOME_STATEMENT.append(p)

                
                #page3
                if "Cash on Hand and in Banks " in line:
                    op_cash = fileData.split("Cash on Hand and in Banks ")[1].split(" ")[1].replace("$","").replace("\nCash-Patient","").replace("1\n2","")
                    cash.append(op_cash)
                    #print(cash)

                if "Cash on Hand and in Banks " in line:
                    ac_cash = fileData.split("Cash on Hand and in Banks ")[1].split(" ")[3].replace("\nCash-Patient","").replace("1\n2","").replace("&","").replace("|","").replace("Cash-Patient","")
                    cash_ac.append(ac_cash)
                    #print(cash_ac)

                if "Cash-Patient Deposits " in line:
                    op_cash_dep = fileData.split("Cash-Patient Deposits ")[1].split(" ")[0].replace("————SSSSSid","").replace("2\nAccounts","")
                    cash_dep.append(op_cash_dep)
                    # print(cash_dep)
                
                if "Cash-Patient Deposits " in line:
                    ac_cash_dep = fileData.split("Cash-Patient Deposits ")[1].split(" ")[1].replace("\n\nAccounts","").replace("Sid","").replace("&","")
                    cash_dep_ac.append(ac_cash_dep)
                    #print(cash_dep_ac)
                

                if "Accounts & Short-Term Notes Receivable3 | Patients (less allowance ) "in line:
                    op_acc = fileData.split("Accounts & Short-Term Notes Receivable3 | Patients (less allowance ) ")[1].split(" ")[0]
                    acc.append(op_acc)
                    #print(acc)
                

                if "Accounts & Short-Term Notes Receivable3 | Patients (less allowance ) "in line:
                    ac_acc = fileData.split("Accounts & Short-Term Notes Receivable3 | Patients (less allowance ) ")[1].split(" ")[1]
                    acc_ac.append(ac_acc)
                    #print(acc_ac)
                

                if "Supply Inventory (priced at ) " in line:
                    op_supply = fileData.split("Supply Inventory (priced at ) ")[1].split(" ")[0].replace("4\n5","")
                    supply_op.append(op_supply)
                    #print(op_supply)
                
                
                if "Supply Inventory (priced at ) " in line:
                    ac_supply = fileData.split("Supply Inventory (priced at ) ")[1].split(" ")[1].replace("|","").replace("4\n5","")
                    supply_ac.append(ac_supply)
                    #print(supply_ac)
                
                
                if "Short-Term Investments " in line:
                    op_short = fileData.split("Short-Term Investments ")[1].split(" ")[0].replace("|","").replace("5\n6","").replace("SSC","")
                    short_op.append(op_short)
                    #print(short_op)
                
                
                if "Short-Term Investments " in line:
                    ac_short = fileData.split("Short-Term Investments ")[1].split(" ")[1].replace("|","").replace("5\n6","").replace("Cd","").replace("CCT","")
                    short_ac.append(ac_short)
                    #print(short_ac)
                
                if "Prepaid Insurance " in line:
                    op_prepaid = fileData.split("Prepaid Insurance ")[1].split(" ")[0]
                    prepaid_op.append(op_prepaid)
                    #print(prepaid_op)
                

                if "Prepaid Insurance " in line:
                    ac_prepaid = fileData.split("Prepaid Insurance ")[1].split(" ")[1]
                    prepaid_ac.append(op_prepaid)
                    #print(prepaid_ac)
                

                if "Other Prepaid Expenses " in line:
                    op_other = fileData.split("Other Prepaid Expenses ")[1].split(" ")[0].replace("7\n8","").replace("SSCs","")
                    other_op.append(op_other)
                    #print(other_op)
                
                
                if "Other Prepaid Expenses " in line:
                    ac_other = fileData.split("Other Prepaid Expenses ")[1].split(" ")[1].replace("7\n8","").replace("|","").replace("SCS","")
                    other_ac.append(ac_other)
                    #print(other_ac)
                

                if "Accounts Receivable (owners or related parties) " in line:
                    op_account = fileData.split("Accounts Receivable (owners or related parties) ")[1].split(" ")[0].replace("8\n9","")
                    account_op.append(op_account)
                    #print(account_op)
                
                if "Accounts Receivable (owners or related parties) " in line:
                    ac_account = fileData.split("Accounts Receivable (owners or related parties) ")[1].split(" ")[1].replace("|","")
                    account_ac.append(ac_account)
                    #print(account_ac)
                

                #bug
                if "9 | Other(specify): " in line:
                    op_other_specify = fileData.split("Other(specify): ")[1].split(" ")[2].replace("See","").replace("Due","")
                    other_specify_op.append(op_other_specify)
                    #print(other_specify_op) 

                if "9 | Other(specify): " in line:
                    ac_other_specify = fileData.split("Other(specify): ")[1].split(" ")[3].replace("CBC\n\nTOTAL","").replace("+i.","").replace("party","").replace("|","").replace("/Escrows","")
                    other_specify_ac.append(ac_other_specify)
                    #print(other_specify_op) 
                

                if "(sum of lines 1 thru 9) " in line:
                    op_current_assets = fileData.split("(sum of lines 1 thru 9) ")[1].split(" ")[1]
                    current_assets_op.append(op_current_assets)
                    #print(current_assets_op)
                
                
                if "(sum of lines 1 thru 9) " in line:
                    ac_current_assets = fileData.split("(sum of lines 1 thru 9) ")[1].split(" ")[3].replace("10\nB.","")
                    current_assets_ac.append(ac_current_assets)
                    #print(current_assets_ac)
            

                if "Long-Term Notes Receivable " in line:
                    op_long = fileData.split("Long-Term Notes Receivable ")[1].split(" ")[0].replace("11\n12","")
                    long_op.append(op_long)
                    #print(long_op)
                
                
                if "Long-Term Notes Receivable " in line:
                    ac_long = fileData.split("Long-Term Notes Receivable ")[1].split(" ")[1].replace("|","")
                    long_ac.append(ac_long)
                    #print(long_ac)
                

                if "Long-Term Investments " in line:
                    op_long_term = fileData.split("Long-Term Investments ")[1].split(" ")[0].replace("12\n13","")
                    long_term_op.append(op_long_term)
                    #print(long_term_op)
                
                if "Long-Term Investments " in line:
                    ac_long_term = fileData.split("Long-Term Investments ")[1].split(" ")[1].replace("12\n13","").replace("|","")
                    long_term_ac.append(ac_long_term)
                    #print(long_term_ac)
            
                
                if "Land " in line:
                    op_land = fileData.split("Land ")[1].split(" ")[0].replace("13\n14","")
                    land_op.append(op_land)
                    #print(land_op)
                
                
                if "Land " in line:
                    ac_land = fileData.split("Land ")[1].split(" ")[1].replace("13\n14","").replace("|","")
                    land_ac.append(ac_land)
                    #print(land_ac)
                

                #bug
                if "Buildings, at Historical Cost " in line:
                    op_build = fileData.split("Buildings, at Historical Cost ")[1].split(" ")[0].replace("14\n15","")
                    build_op.append(op_build)
                    #print(build_op)
                

                #bug
                if "Buildings, at Historical Cost " in line:
                    ac_build = fileData.split("Buildings, at Historical Cost ")[1].split(" ")[0]
                    build_ac.append(ac_build)
                    #print(build_ac)
                

                if "Leasehold Improvements, at Historical Cost " in line:
                    op_leasehold = fileData.split("Leasehold Improvements, at Historical Cost ")[1].split(" ")[0]
                    leasehold_op.append(op_leasehold)
                    #print(leasehold_op)
                

                if "Leasehold Improvements, at Historical Cost " in line:
                    ac_leasehold = fileData.split("Leasehold Improvements, at Historical Cost ")[1].split(" ")[1]
                    leasehold_ac.append(ac_leasehold)
                    #print(leasehold_ac)
                
                
                if "Equipment, at Historical Cost " in line:
                    op_equipment = fileData.split("Equipment, at Historical Cost ")[1].split(" ")[0]
                    equipment_op.append(op_equipment)
                    #print(equipment_op)
                
                
                if "Equipment, at Historical Cost " in line:
                    ac_equipment = fileData.split("Equipment, at Historical Cost ")[1].split(" ")[1]
                    equipment_ac.append(ac_equipment)
                    #print(equipment_ac)
                

                if "Accumulated Depreciation (book methods) " in line:
                    op_accumulated = fileData.split("Accumulated Depreciation (book methods) ")[1].split(" ")[0].replace("(","").replace(")","")
                    accumulated_op.append(op_accumulated)
                    #print(accumulated_op)
                

                if "Accumulated Depreciation (book methods) " in line:
                    ac_accumulated = fileData.split("Accumulated Depreciation (book methods) ")[1].split(" ")[1].replace("(","").replace(")","").replace("17\n18","")
                    accumulated_ac.append(ac_accumulated)
                    #print(accumulated_ac)
                

                if "Deferred Charges " in line:
                    op_deferred = fileData.split("Deferred Charges ")[1].split(" ")[0].replace("18\n19","")
                    deferred_op.append(op_deferred)
                    #print(deferred_op)
                

                if "Deferred Charges " in line:
                    ac_deferred = fileData.split("Deferred Charges ")[1].split(" ")[1].replace("|","")
                    deferred_ac.append(ac_deferred)
                    #print(deferred_ac)
                
                #bug
                if "19 | Organization & Pre-Operating Costs " in line:
                    op_organization = fileData.split("Organization & Pre-Operating Costs ")[1].split(" ")[0].replace("19\nAccumulated","")
                    organization_op.append(op_organization)
                    #print(organization_op)
                

                #bug
                if "19 | Organization & Pre-Operating Costs " in line:
                    ac_organization = fileData.split("Organization & Pre-Operating Costs ")[1].split(" ")[1].replace("Amortization","")
                    organization_ac.append(ac_organization)
                    #print(organization_ac)
                

                if "Accumulated Amortization 20 | Organization & Pre-Operating Costs " in line:
                    op_accumulated_amortization = fileData.split("Accumulated Amortization 20 | Organization & Pre-Operating Costs ")[1].split(" ")[0].replace("20\n21","").replace("(","").replace(")","")
                    accumulated_amortization_op.append(op_accumulated_amortization)
                    #print(accumulated_amortization_op)
                
                
                if "Accumulated Amortization 20 | Organization & Pre-Operating Costs " in line:
                    ac_accumulated_amortization = fileData.split("Accumulated Amortization 20 | Organization & Pre-Operating Costs ")[1].split(" ")[1].replace("20\n21","").replace("|","").replace(")","")
                    accumulated_amortization_ac.append(ac_accumulated_amortization)
                    #print(accumulated_amortization_ac)
                

                if "Restricted Funds " in line:
                    op_restricted = fileData.split("Restricted Funds ")[1].split(" ")[0].replace("21\n22","").replace("|\nOther","")
                    restricted_op.append(op_restricted)
                    #print(restricted_op)
                

                if "Restricted Funds " in line:
                    ac_restricted = fileData.split("Restricted Funds ")[1].split(" ")[1].replace("21\n22","").replace("Long-Term","").replace("|","")
                    restricted_ac.append(ac_restricted)
                    #print(restricted_ac)
                
                if "Other Long-Term Assets (specify): " in line:
                    op_other_long_term = fileData.split("Other Long-Term Assets (specify): ")[1].split(" ")[0].replace("22\n23","").replace("|\nTOTAL","").replace("|\n|\n$","")
                    other_long_term_op.append(op_other_long_term)
                    #print(other_long_term_op)
                

                if "Other Long-Term Assets (specify): " in line:
                    ac_other_long_term = fileData.split("Other Long-Term Assets (specify): ")[1].split(" ")[1].replace("Long-Term","").replace("2\n\nTOTAL","").replace("22\n23","").replace("|","")
                    other_long_term_ac.append(ac_other_long_term)
                    #print(other_long_term_ac)
                

                #bug
                if "23 | Other(specify): " in line:
                    op_other_specify_two = fileData.split("23 | Other(specify): ")[1].split(" ")[0].replace("See","").replace("Due","")
                    other_specify_two_op.append(op_other_specify_two)
                    #print(other_specify_two_op)
                

                #bug
                if "23 | Other(specify): " in line:
                    ac_other_specify_two = fileData.split("23 | Other(specify): ")[1].split(" ")[3].replace("|","").replace("&","").replace("Assets","")
                    other_specify_two_ac.append(ac_other_specify_two)
                    #print(other_specify_two_ac)
            
                if "(sum of lines 11 thru 23) " in line:
                    op_long_term_assets = fileData.split("(sum of lines 11 thru 23) ")[1].split(" ")[1].replace("|$","")
                    long_term_assets_op.append(op_long_term_assets)
                    #print(long_term_assets_op)
                

                if "(sum of lines 11 thru 23) " in line:
                    ac_long_term_assets = fileData.split("(sum of lines 11 thru 23) ")[1].split(" ")[3].replace("24\nTOTAL","").replace("ASSETS\n25","").replace("|","")
                    long_term_assets_ac.append(ac_long_term_assets)
                    #print(long_term_assets_ac)
            
                if "(sum of lines 10 and 24) " in line:
                    op_total_assets = fileData.split("(sum of lines 10 and 24) ")[1].split(" ")[1]
                    total_assets_op.append(op_total_assets)
                    #print(total_assets_op)
                

                if "(sum of lines 10 and 24) " in line:
                    ac_total_assets = fileData.split("(sum of lines 10 and 24) ")[1].split(" ")[3].replace("25\n\nHFS","").replace("|$","")
                    total_assets_ac.append(ac_total_assets)
                    #print(total_assets_ac)
                
                
                if "26 | Accounts Payable " in line:
                    op_accounts_payable = fileData.split("26 | Accounts Payable ")[1].split(" ")[1]
                    accounts_payable_op.append(op_accounts_payable)
                    #print(accounts_payable_op)
                

                if "26 | Accounts Payable " in line:
                    ac_accounts_payable = fileData.split("26 | Accounts Payable ")[1].split(" ")[3].replace("|","")
                    accounts_payable_ac.append(ac_accounts_payable)
                    #print(accounts_payable_ac)
                

                if "Officer's Accounts Payable " in line:
                    op_offer = fileData.split("Officer's Accounts Payable ")[1].split(" ")[0].replace("27\n28","").replace("———d|","")
                    offer_op.append(op_offer)
                    #print(offer_op)
                
                
                if "Officer's Accounts Payable " in line:
                    ac_offer = fileData.split("Officer's Accounts Payable ")[1].split(" ")[1].replace("|","").replace("—SSSSSCSC~id","")
                    offer_ac.append(ac_offer)
                    #print(offer_ac)
                
                
                if "Accounts Payable-Patient Deposits " in line:
                    op_accounts_patient = fileData.split("Accounts Payable-Patient Deposits ")[1].split(" ")[0].replace("[+4","").replace("28\n29","")
                    accounts_patient_op.append(op_accounts_patient)
                    #print(accounts_patient_op)
                

                if "Accounts Payable-Patient Deposits " in line:
                    ac_accounts_patient = fileData.split("Accounts Payable-Patient Deposits ")[1].split(" ")[1].replace("SS","").replace("|","").replace("28\n29","")
                    accounts_patient_ac.append(ac_accounts_patient)
                    #print(accounts_patient_ac)
                

                if "Short-Term Notes Payable " in line:
                    op_sort_term = fileData.split("Short-Term Notes Payable ")[1].split(" ")[0].replace("—————s[SSSSidT","")
                    sort_term_op.append(op_sort_term)
                    #print(sort_term_op)
                
                
                if "Short-Term Notes Payable " in line:
                    ac_sort_term = fileData.split("Short-Term Notes Payable ")[1].split(" ")[1].replace("SSSCS—~d","").replace("29\n30","")
                    sort_term_ac.append(ac_sort_term)
                    #print(sort_term_ac)
                

                if "Accrued Salaries Payable " in line:
                    op_accured_salaries = fileData.split("Accrued Salaries Payable ")[1].split(" ")[0].replace("__——~+d|","").replace("T7537","")
                    accured_salaries_op.append(op_accured_salaries)
                    #print(accured_salaries_op)
                

                if "Accrued Salaries Payable " in line:
                    ac_accured_salaries = fileData.split("Accrued Salaries Payable ")[1].split(" ")[1].replace("=i","").replace("L337","").replace("30\nAccrued","")
                    accured_salaries_ac.append(ac_accured_salaries)
                    #print(accured_salaries_ac)
                

                if "(excluding real estate taxes) " in line:
                    op_accrued_taxes = fileData.split("(excluding real estate taxes) ")[1].split(" ")[0].replace("3\n[7_[","")
                    accrued_taxes_op.append(op_accrued_taxes)
                    #print(accrued_taxes_op)
            

                if "(excluding real estate taxes) " in line:
                    ac_accrued_taxes = fileData.split("(excluding real estate taxes) ")[1].split(" ")[1].replace("31\n32","").replace("Other","")
                    accrued_taxes_ac.append(ac_accrued_taxes)
                    #print(accrued_taxes_ac)
                

                if "Accrued Real Estate Taxes(Sch.[X-B) " in line:
                    op_accrued_real = fileData.split("Accrued Real Estate Taxes(Sch.[X-B) ")[1].split(" ")[0]
                    accrued_real_op.append(op_accrued_real)
                    #print(accrued_real_op)
                

                if "Accrued Real Estate Taxes(Sch.[X-B) " in line:
                    ac_accrued_real = fileData.split("Accrued Real Estate Taxes(Sch.[X-B) ")[1].split(" ")[1].replace("32\n33","")
                    accrued_real_ac.append(ac_accrued_real)
                    #print(accrued_real_ac)
                

                if "Accrued Interest Payable " in line:
                    op_accrued_interest = fileData.split("Accrued Interest Payable ")[1].split(" ")[0].replace("33\n34","").replace("S|","")
                    accrued_interest_op.append(op_accrued_interest)
                    #print(accrued_interest_op)
                

                if "Accrued Interest Payable " in line:
                    ac_accrued_interest = fileData.split("Accrued Interest Payable ")[1].split(" ")[1].replace("33\n34","").replace("|","").replace("Cd\n4","")
                    accrued_interest_ac.append(ac_accrued_interest)
                    #print(accrued_interest_ac)
                

                if "34 | Deferred Compensation " in line:
                    op_deferred = fileData.split("34 | Deferred Compensation ")[1].split(" ")[0].replace("34\n35","")
                    deferred_compensation_op.append(op_deferred)
                    #print(deferred_op)
                
                
                if "34 | Deferred Compensation " in line:
                    ac_deferred = fileData.split("34 | Deferred Compensation ")[1].split(" ")[1].replace("|","")
                    deferred_compensation_ac.append(ac_deferred)
                    #print(deferred_ac)
                

                if "Federal and State Income Taxes " in line:
                    op_federal = fileData.split("Federal and State Income Taxes ")[1].split(" ")[0].replace("35\nOther","").replace("|","")
                    federal_op.append(op_federal)
                    #print(federal_op)
                
                
                if "Federal and State Income Taxes " in line:
                    ac_federal = fileData.split("Federal and State Income Taxes ")[1].split(" ")[1].replace("|\nae\n\nTOTAL","").replace("Current","")
                    federal_ac.append(ac_federal)
                    #print(federal_ac)
                
                
                if "Due to Partnership " in line:
                    op_current_liabilities = fileData.split("Due to Partnership ")[1].split(" ")[0]
                    current_liabilities_op.append(op_current_liabilities)
                    #print(current_liabilities_op)
                
                
                if "Due to Partnership " in line:
                    ac_current_liabilities = fileData.split("Due to Partnership ")[1].split(" ")[1]
                    current_liabilities_ac.append(ac_current_liabilities)
                    #print(current_liabilities_ac)
                
                
                if "(sum of lines 26 thru 37) " in line:
                    op_total_current_liabilities = fileData.split("(sum of lines 26 thru 37) ")[1].split(" ")[1].replace("\nD.","")
                    total_current_liabilities_op.append(op_total_current_liabilities)
                    #print(total_current_liabilities_op)
                
                if "(sum of lines 26 thru 37) " in line:
                    ac_total_current_liabilities = fileData.split("(sum of lines 26 thru 37) ")[1].split(" ")[3].replace("\nD.","").replace("Liabilities\n\n6\n\n38\n[","").replace("Long-Term","")
                    total_current_liabilities_ac.append(ac_total_current_liabilities)
                    #print(total_current_liabilities_ac)
                

                if "39 | Long-Term Notes Payable " in line:
                    op_long_term_notes = fileData.split("39 | Long-Term Notes Payable ")[1].split(" ")[0].replace("39\n40","")
                    long_term_notes_op.append(op_long_term_notes)
                    #print(long_term_notes_op)
                

                if "39 | Long-Term Notes Payable " in line:
                    ac_long_term_notes = fileData.split("39 | Long-Term Notes Payable ")[1].split(" ")[1].replace("39\n40","").replace("|","")
                    long_term_notes_ac.append(ac_long_term_notes)
                    #print(long_term_notes_ac)
                
                
                if "Mortgage Payable " in line:
                    op_mortgage = fileData.split("Mortgage Payable ")[1].split(" ")[0].replace("40\n41","").replace("ee","")
                    mortgage_op.append(op_mortgage)
                    #print(mortgage_op)
                
                if "Mortgage Payable " in line:
                    ac_mortgage = fileData.split("Mortgage Payable ")[1].split(" ")[1].replace("40\n41","").replace("ee\nBonds","").replace("|","")
                    mortgage_ac.append(ac_mortgage)
                    #print(mortgage_ac)
                

                if "Bonds Payable " in line:
                    op_bond = fileData.split("Bonds Payable ")[1].split(" ")[0].replace("41\n42","").replace("Sd\n[Deferred","").replace("re\nDeferred","")
                    bond_op.append(op_bond)
                    #print(bond_op)
                

                if "Bonds Payable " in line:
                    ac_bond = fileData.split("Bonds Payable ")[1].split(" ")[1].replace("Compensation","").replace("|","")
                    bond_ac.append(ac_bond)
                    #print(bond_ac)
                
                
                if "42 | Deferred Compensation " in line:
                    op_deferred_compensation = fileData.split("42 | Deferred Compensation ")[1].split(" ")[0].replace("42\nOther","")
                    deferred_compensation_two_op.append(op_deferred_compensation)
                    #print(deferred_compensation_op)
                

                if "(sum of lines 39 thru 44) " in line:
                    op_long_liabilities = fileData.split("(sum of lines 39 thru 44) ")[1].split(" ")[1].replace("\nTOTAL","")
                    long_liabilities_op.append(op_long_liabilities)
                    #print(long_liabilities_op)
                
                
                if "(sum of lines 39 thru 44) " in line:
                    ac_long_liabilities = fileData.split("(sum of lines 39 thru 44) ")[1].split(" ")[3].replace("45\nTOTAL","").replace("TOTAL","").replace("|$\n\n","").replace("LIABILITIES\n46","")
                    long_liabilities_ac.append(ac_long_liabilities)
                    #print(long_liabilities_ac)
                

                if "(sum of lines 38 and 45) " in line:
                    op_total_liabilities = fileData.split("(sum of lines 38 and 45) ")[1].split(" ")[1].replace("\n7","")
                    total_liabilities_op.append(op_total_liabilities)
                    #print(total_liabilities_op)
                

                if "(sum of lines 38 and 45) " in line:
                    ac_total_liabilities = fileData.split("(sum of lines 38 and 45) ")[1].split(" ")[3].replace("\n\npo\nre","").replace("TOTAL","").replace("|$","").replace("|","").replace("\n7","")
                    total_liabilities_ac.append(ac_total_liabilities)
                    #rint(total_liabilities_ac)
            
                if "TOTAL EQUITY (page 18, line 24) " in line:
                    op_total_equity = fileData.split("TOTAL EQUITY (page 18, line 24) ")[1].split(" ")[1].replace("LIABILITIES","").replace("(","").replace(")","")
                    total_equity_op.append(op_total_equity)
                    #print(total_equity_op)
            
                
                if "TOTAL EQUITY (page 18, line 24) " in line:
                    ac_total_equity = fileData.split("TOTAL EQUITY (page 18, line 24) ")[1].split(" ")[3].replace("LIABILITIES","").replace("(","").replace(")","").replace("|$","").replace("EQUITY\nsum","")
                    total_equity_ac.append(ac_total_equity)
                    #print(total_equity_ac)
                

                if "(sum of lines 46 and 47) " in line:
                    op_liabilities_equity = fileData.split("(sum of lines 46 and 47) ")[1].split(" ")[1].replace("\n\na»\n\na»\n\ni","")
                    liabilities_equity_op.append(op_liabilities_equity)
                    #print(liabilities_equity_op)
                
                if "(sum of lines 46 and 47) " in line:
                    ac_liabilities_equity = fileData.split("(sum of lines 46 and 47) ")[1].split(" ")[3].replace("|$","").replace("ES","").replace("Name","")
                    liabilities_equity_ac.append(ac_liabilities_equity)
                    #print(liabilities_equity_ac)
                
                
                # page 4
                if "Gross Revenue -- All Levels of Care " in line:
                    gross_amount = (
                        fileData.split("Gross Revenue -- All Levels of Care ")[1]
                        .split(" ")[1]
                        .replace("1\n2", "")
                    )
                    gross.append(gross_amount)
                    # print(gross)

                if "Discounts and Allowances for all Levels " in line:
                    dis_all = (
                        fileData.split("Discounts and Allowances for all Levels ")[1]
                        .split(" ")[0]
                        .replace("(", "")
                        .replace("2\n3", "")
                    )
                    disc.append(dis_all)
                    #print(disc)

                if "SUBTOTAL Inpatient Care (line 1 minus line 2) " in line:
                    sub = (
                        fileData.split("SUBTOTAL Inpatient Care (line 1 minus line 2) ")[1]
                        .split(" ")[0]
                        .replace("$", "")
                    )
                    sub_total.append(sub)
                    # print(sub_total)

                if "Day Care " in line:
                    day = (
                        fileData.split("Day Care ")[1]
                        .split(" ")[0]
                        .replace("4\n5", "")
                        .replace("4", "")
                    )
                    day_care.append(day)
                    # print(day_care)

                if "Other Care for Outpatients " in line:
                    other = (
                        fileData.split("Other Care for Outpatients ")[1]
                        .split(" ")[0]
                        .replace("5\n6", "")
                    )
                    other_care.append(other)
                    # print(other_care)

                if "Therapy " in line:
                    the = fileData.split("Therapy ")[1].split(" ")[0].replace("6\n7", "")
                    therapy.append(the)
                    # print(therapy)

                if "Oxygen " in line:
                    oxy = fileData.split("Oxygen ")[1].split(" ")[0].replace("i\n8", "").replace("7\n8", "")
                    oxygen.append(oxy)
                    #print(oxygen)
                
                if "SUBTOTAL Ancillary Revenue (lines 4 thru 7) " in line:
                    sub = fileData.split("SUBTOTAL Ancillary Revenue (lines 4 thru 7) ")[1].split(" ")[0].replace("8\n[EIC.","").replace("$","")
                    subtotal_ancillary.append(sub)
                    #print(subtotal_ancillary)

                if "Payments for Education " in line:
                    pay = fileData.split("Payments for Education ")[1].split(" ")[0].replace("9\n10","")
                    payment.append(pay)
                    #print(payment)
                
                if "Other Government Grants " in line:
                    other = fileData.split("Other Government Grants ")[1].split(" ")[0].replace("10\n11","")
                    other_gov.append(other)
                    #print(other_gov)
                
                if "CNA Training Reimbursements " in line:
                    cna = fileData.split("CNA Training Reimbursements ")[1].split(" ")[0].replace("11\nTZ","").replace("11\n12","")
                    can.append(cna)
                    #print(can)

                if "Gift and Coffee Shop " in line:
                    gift_coffee = fileData.split("Gift and Coffee Shop ")[1].split(" ")[0].replace("IZ\n13","").replace("12\n13","").replace("12","")
                    gift.append(gift_coffee)
                    #print(gift)

                if "Barber and Beauty Care " in line:
                    bar = fileData.split("Barber and Beauty Care ")[1].split(" ")[0].replace("13\n14","")
                    barber.append(bar)
                    #print(barber)
                
                if "Non-Patient Meals " in line:
                    non = fileData.split("Non-Patient Meals ")[1].split(" ")[0].replace("14\n15","")
                    non_pat.append(non)
                    #print(non_pat)

                if "Telephone, Television and Radio " in line:
                    tel = fileData.split("Telephone, Television and Radio ")[1].split(" ")[0].replace("15\n16","")
                    tele.append(tel)
                    #print(tele)
                
                if "Rental of Facility Space " in line:
                    rent = fileData.split("Rental of Facility Space ")[1].split(" ")[0].replace("16\n17","")
                    rental.append(rent)
                    #print(rental)
                
                if "Sale of Drugs " in line:
                    drug = fileData.split("Sale of Drugs ")[1].split(" ")[0].replace("17\nT8","")
                    sale_drug.append(drug)
                    #print(sale_drug)
                
                if "Sale of Supplies to Non-Patients " in line:
                    sup = fileData.split("Sale of Supplies to Non-Patients ")[1].split(" ")[0].replace("18\n19","")
                    sale_sup.append(sup)
                    #print(sale_sup)
                
                if "Laboratory " in line:
                    laboratory = fileData.split("Laboratory ")[1].split(" ")[0].replace("19\n20","")
                    lab.append(laboratory)
                    #print(lab)

                if "Radiology and X-Ray " in line:
                    radiology = fileData.split("Radiology and X-Ray ")[1].split(" ")[0].replace("20\n21","")
                    radio.append(radiology)
                    #print(radio)
                
                if "Other Medical Services " in line:
                    other_medical = fileData.split("Other Medical Services ")[1].split(" ")[0].replace("21\n22","")
                    other_medi.append(other_medical)
                    #print(other_medi)
                
                if "Laundry " in line:
                    laundri = fileData.split("Laundry ")[1].split(" ")[0].replace("22\n23","")
                    laundry.append(laundri)
                    #print(laundry)
                
                #bug
                if "SUBTOTAL Other Operating Revenue (lines 9 thru 22 " in line:
                    subtotal_ope = fileData.split("SUBTOTAL Other Operating Revenue (lines 9 thru 22 ")[1].split(" ")[0].replace("23\nD.","")
                    subtotal_other.append(subtotal_ope)
                    #print(subtotal_other)

                if "Contributions " in line:
                    contributions = fileData.split("Contributions ")[1].split(" ")[0].replace("24\n25","")
                    contri.append(contributions)
                    #print(contri)

                if "Interest and Other Investment Income*** " in line:
                    inter = fileData.split("Interest and Other Investment Income*** ")[1].split(" ")[0].replace("25\n26","")
                    interest.append(inter)
                    #print(interest)
                
                if "SUBTOTAL Non-Operating Revenue (lines 24 and 25) " in line:
                    subtotal_non = fileData.split("SUBTOTAL Non-Operating Revenue (lines 24 and 25) ")[1].split(" ")[0].replace("|$","")
                    subtotal_non_op.append(subtotal_non)
                    #print(subtotal_non_op)
                
                if "Settlement Income (Insurance, Legal, Etc.) " in line:
                    settlement_income = fileData.split("Settlement Income (Insurance, Legal, Etc.) ")[1].split(" ")[0].replace("27\n28","")
                    settlement.append(settlement_income)
                    #print(settlement)

                # if "Government Stimulus Income " in line:
                #     government = fileData.split("Government Stimulus Income ")[1].split(" ")[0]
                #     gov.append(government)
                #     print(gov)

                if "SUBTOTAL Other Revenue (lines 27, 28 and 28a) " in line:
                    subtotal_revenue = fileData.split("SUBTOTAL Other Revenue (lines 27, 28 and 28a) ")[1].split(" ")[0].replace("$","")
                    subtotal_other_revenue.append(subtotal_revenue)
                    #print(subtotal_other_revenue)

                if "TOTAL REVENUE (sum of lines 3, 8, 23, 26 and 29) " in line:
                    revenue = fileData.split("TOTAL REVENUE (sum of lines 3, 8, 23, 26 and 29) ")[1].split(" ")[0].replace("|$","").replace("_","").replace("—","")
                    total_revenue.append(revenue)
                    #print(total_revenue)

                #bug
                if "General Services " in line:
                    services = fileData.split("General Services ")[1].split(" ")[0].replace("\nHealth","").replace("\n\n32","").replace("\n2","")
                    general.append(services)
                    #print(general)

                #bug
                if "Health Care " in line:
                    health_care = fileData.split("Health Care ")[1].split(" ")[0].replace("\n3","")
                    health.append(health_care)
                    #print(health)
                
                #bug
                if "General Administration " in line:
                    admin = fileData.split("General Administration ")[1].split(" ")[0].replace("\nB.","").replace("\n[EB","").replace(".","")
                    general_admin.append(admin)
                    #print(general_admin)
                
                #bug
                if "Ownership " in line:
                    own = fileData.split("Ownership ")[1].split(" ")[0].replace("\n\nC","").replace("\n\nProvider","").replace("\n5","").replace("\n\nSpecial","").replace(".","")
                    ownership.append(own)
                    #print(ownership)

                #bug
                if "Special Cost Centers " in line:
                    cost = fileData.split("Special Cost Centers ")[1].split(" ")[0]
                    special_cost.append(cost)
                    #print(special_cost)

                #bug
                if "Provider Participation Fee " in line:
                    provider_fee = fileData.split("Provider Participation Fee ")[1].split(" ")[0]
                    provider.append(provider_fee)
                    #print(provider)

                if "TOTAL EXPENSES (sum of lines 31 thru 39)* " in line:
                    expenses = fileData.split("TOTAL EXPENSES (sum of lines 31 thru 39)* ")[1].split(" ")[0].replace("\n41","").replace("\n13","").replace("$","")
                    total_expenses.append(expenses)
                    #print(total_expenses)

                if "Income before Income Taxes (line 30 minus line 40)** " in line:
                    before = fileData.split("Income before Income Taxes (line 30 minus line 40)** ")[1].split(" ")[0].replace("(","").replace(")","")
                    income_before.append(before)
                    #print(income_before)

                #bug
                if "42 |Income Taxes " in line:
                    taxes = fileData.split("42 |Income Taxes ")[1].split(" ")[0].replace("42\n\n43","").replace("42\n43","").replace("42\n17","")
                    income_tax.append(taxes)
                    #print(income_tax)

                if "NET INCOME OR LOSS FOR THE YEAR (line 41 minus line 42) |$ " in line:
                    net = fileData.split("NET INCOME OR LOSS FOR THE YEAR (line 41 minus line 42) |$ ")[1].split(" ")[0].replace("(","").replace(")","")
                    net_income.append(net)
                    #print(net_income)
                
                if "Medicaid - Net Inpatient Revenue " in line:
                    medi = fileData.split("Medicaid - Net Inpatient Revenue ")[1].split(" ")[1]
                    madicaid.append(medi)
                    #print(madicaid)
                
                if "Private Pay - Net Inpatient Revenue " in line:
                    private_pay = fileData.split("Private Pay - Net Inpatient Revenue ")[1].split(" ")[0].replace("45\n23","").replace("45\nMedicare","")
                    private.append(private_pay)
                    #print(private)
                
            
                if "Medicare - Net Inpatient Revenue " in line:
                    medi_care = fileData.split("Medicare - Net Inpatient Revenue ")[1].split(" ")[0].replace("46\nOther-(specify)","").replace("46\nD","")
                    medicare.append(medi_care)
                    #print(medicare)

                #bug
                if other_specific_counter == 1:
                    if "Other-(specify) " in line:
                        med_b = line.split("Other-(specify) ")[1].split(" ")
                        if len(med_b) >= 2:
                            other_specify_med.append(med_b[-2].replace("Hospice","").replace("(","").replace(")",""))
                        other_specific_counter = 2
                        #print(other_specify_med)
                else:
                    #bug
                    if "Other-(specify) " in line:
                        specify = line.split("Other-(specify) ")[1].split(" ")
                        if len(specify) >= 2:
                            other_speci.append(specify[-2].replace("Allows","").replace("(","").replace(")",""))
                        
                        #print(other_specify)

                if "TOTAL Inpatient Care Revenue (This total must agree to Line 3) " in line:
                    inpa = fileData.split("TOTAL Inpatient Care Revenue (This total must agree to Line 3) ")[1].split(" ")[1]
                    total_inpa.append(inpa)
                    #print(total_inpa)

            cash = check_list(cash, file_counter)
            cash_ac = check_list(cash_ac, file_counter)
            cash_dep = check_list(cash_dep, file_counter)
            cash_dep_ac = check_list(cash_dep_ac, file_counter)
            acc = check_list(acc, file_counter)
            acc_ac = check_list(acc_ac, file_counter)
            supply_op = check_list(supply_op, file_counter)
            supply_ac = check_list(supply_ac, file_counter)
            short_op = check_list(short_op, file_counter)
            short_ac = check_list(short_ac, file_counter)
            prepaid_op = check_list(prepaid_op, file_counter)
            prepaid_ac = check_list(prepaid_ac, file_counter)
            other_op = check_list(other_op, file_counter)
            other_ac = check_list(other_ac, file_counter)
            account_op = check_list(account_op, file_counter)
            account_ac = check_list(account_ac, file_counter)
            other_specify_op = check_list(other_specify_op, file_counter)
            other_specify_ac = check_list(other_specify_ac, file_counter)
            current_assets_op = check_list(current_assets_op, file_counter)
            current_assets_ac = check_list(current_assets_ac, file_counter)
            long_op = check_list(long_op, file_counter)
            long_ac = check_list(long_ac, file_counter)
            long_term_op = check_list(long_term_op, file_counter)
            long_term_ac = check_list(long_term_ac, file_counter)
            land_op = check_list(land_op, file_counter)
            land_ac = check_list(land_ac, file_counter)
            build_op = check_list(build_op, file_counter)
            build_ac = check_list(build_ac, file_counter)
            leasehold_op = check_list(leasehold_op, file_counter)
            leasehold_ac = check_list(leasehold_ac, file_counter)
            equipment_op = check_list(equipment_op, file_counter)
            equipment_ac = check_list(equipment_ac, file_counter)
            accumulated_op = check_list(accumulated_op, file_counter)
            accumulated_ac = check_list(accumulated_ac, file_counter)
            deferred_op = check_list(deferred_op, file_counter)
            deferred_ac = check_list(deferred_ac, file_counter)
            organization_op = check_list(organization_op, file_counter)
            organization_ac = check_list(organization_ac, file_counter)
            accumulated_amortization_op = check_list(accumulated_amortization_op, file_counter)
            accumulated_amortization_ac = check_list(accumulated_amortization_ac, file_counter)
            restricted_op = check_list(restricted_op, file_counter)
            restricted_ac = check_list(restricted_ac , file_counter)
            other_long_term_op = check_list(other_long_term_op, file_counter)
            other_long_term_ac = check_list(other_long_term_ac, file_counter)
            other_specify_two_op = check_list(other_specify_two_op, file_counter)
            other_specify_two_ac = check_list(other_specify_two_ac, file_counter)
            long_term_assets_op = check_list(long_term_assets_op, file_counter)
            long_term_assets_ac = check_list(long_term_assets_ac, file_counter)
            total_assets_op = check_list(total_assets_op, file_counter)
            total_assets_ac = check_list(total_assets_ac, file_counter)
            accounts_payable_op = check_list(accounts_payable_op, file_counter)
            accounts_payable_ac = check_list(accounts_payable_ac, file_counter)
            offer_op = check_list(offer_op, file_counter)
            offer_ac = check_list(offer_ac, file_counter)
            accounts_patient_op = check_list(accounts_patient_op, file_counter)
            accounts_patient_ac = check_list(accounts_patient_ac, file_counter)
            sort_term_op = check_list(sort_term_op, file_counter)
            sort_term_ac = check_list(sort_term_ac, file_counter)
            accured_salaries_op = check_list(accured_salaries_op, file_counter)
            accured_salaries_ac = check_list(accured_salaries_ac, file_counter)
            accrued_taxes_op = check_list(accrued_taxes_op, file_counter)
            accrued_taxes_ac = check_list(accrued_taxes_ac, file_counter)
            accrued_real_op = check_list(accrued_real_op, file_counter)
            accrued_real_ac = check_list(accrued_real_ac, file_counter)
            accrued_interest_op = check_list(accrued_interest_op, file_counter)
            accrued_interest_ac = check_list(accrued_interest_ac, file_counter)
            deferred_compensation_op = check_list(deferred_compensation_op, file_counter)
            deferred_compensation_ac = check_list(deferred_compensation_ac, file_counter)
            federal_op = check_list(federal_op, file_counter)
            federal_ac = check_list(federal_ac, file_counter)
            current_liabilities_op = check_list(current_liabilities_op, file_counter)
            current_liabilities_ac = check_list(current_liabilities_ac, file_counter)
            long_term_notes_op = check_list(long_term_notes_op, file_counter)
            long_term_notes_ac = check_list(long_term_notes_ac, file_counter)
            mortgage_op = check_list(mortgage_op, file_counter)
            mortgage_ac = check_list(mortgage_ac, file_counter)
            bond_op = check_list(bond_op, file_counter)
            bond_ac = check_list(bond_ac, file_counter)
            deferred_compensation_two_op = check_list(deferred_compensation_two_op, file_counter)
            deferred_compensation_two_ac = check_list(deferred_compensation_two_ac, file_counter)
            long_liabilities_op = check_list(long_liabilities_op, file_counter)
            long_liabilities_ac = check_list(long_liabilities_ac, file_counter)
            total_liabilities_op = check_list(total_liabilities_op, file_counter)
            total_liabilities_ac = check_list(total_liabilities_ac, file_counter)
            total_equity_op = check_list(total_equity_op, file_counter)
            total_equity_ac = check_list(total_equity_ac, file_counter)
            liabilities_equity_op = check_list(liabilities_equity_op, file_counter)
            liabilities_equity_ac = check_list(liabilities_equity_ac, file_counter)
            subtotal_other = check_list(subtotal_other,file_counter)
            general = check_list(general,file_counter)
            health = check_list(health,file_counter)
            general_admin = check_list(general_admin,file_counter)
            ownership = check_list(ownership,file_counter)
            special_cost = check_list(special_cost,file_counter)
            provider = check_list(provider,file_counter)
            income_tax = check_list(income_tax,file_counter)
            other_specify_med = check_list(other_specify_med,file_counter)
            other_speci = check_list(other_speci,file_counter)
            name = check_list(name,file_counter)
            title = check_list(title,file_counter)
            Firm_Name = check_list(Firm_Name,file_counter)
            address1 = check_list(address1,file_counter)

            file_counter += 1   
            
            
        # print name and title
        name_title = []
        for x in range(len(title)):
            i = str(name[x]) + ", " + str(title[x])
            name_title.append(i)

        # # Firm name and address
        name_address = []
        for y in range(len(address1)):
            l = str(Firm_Name[y]) + ", " + str(address1[y])
            name_address.append(l)
        out_doc = []
        for file in allFiles:
            print("Converting Pdf File To Word: " + file[:-4])
            doc = aw.Document("documents/" + file)

            # convert PDF to Word DOCX format
            doc.save("out_text/" + file[:-4] + "_" + str(date_name) + ".docx")
            out_doc.append(file[:-4] + "_" + str(date_name) + ".docx")
            print("Converted Word: " + file[:-4])
        
        for i in out_doc:
            wordDoc = Document("out_text/" + str(i))

            count = 0
            files = 0

            for table in wordDoc.tables:
                files = 1
                for row in table.rows:
                    row_cells = [_Cell(tc, table) for tc in row._tr.tc_lst]
                    max_cell = 0
                    for cell in row_cells:
                        max_cell += 1
                        val = cell.text.strip().replace("\n", " ")
                        val = re.sub(" +", " ", val)

                        if val == "Beds at Beginning of Report Period":
                            count = 1
                        if count == 1:
                            he = val
                            header.append(he)

                        if count > 1 and count < 8 and max_cell < 6:
                            ro = val
                            listrow.append(ro)
                    if count == 1 or count > 1:
                        count += 1
            if files == 0:
                for x in range(30):
                    listrow.append("")
        
        header.pop()
        tableheader = []
        for i in range(1, 7):
            for j in header[:4]:
                tableheader.append(str(j) + "-" + str(i))

        for min in range(len(allFiles) * 5):
            listrow.pop(min * 5 - min)
        #print(listrow)

        for i in range(len(listrow)):
            if listrow[i] == []:
                listrow[i] = [0]

        allValues = []

        for i in range(24):
            ct = i
            temp = []
            for j in range(len(allFiles)):
                temp.append(listrow[ct])
                ct += 24
            allValues.append(temp)

        dict1 = {
            "IDPH License ID Number": LicenceID,
            "Facility Name": Facility_Name,
            "Address": address,
            "County": country,
            "date": date,
            "Telephone Number": Telephone_Number,
            "print name & title": name_title,
            "Firm Name & Address": name_address,
            "Telephone": Telephone,
            "Percent Occupancy": Percent,
            "BALANCE SHEET": Unrestricted,
            "Beds at Beginning of Report Period-1": allValues[0],
            "Licensure Level of Care-1": allValues[1],
            "Beds at End of Report Period-1": allValues[2],
            "Licensed Bed Days During Report Period-1": allValues[3],
            "Beds at Beginning of Report Period-2": allValues[4],
            "Licensure Level of Care-2": allValues[5],
            "Beds at End of Report Period-2": allValues[6],
            "Licensed Bed Days During Report Period-2": allValues[7],
            "Beds at Beginning of Report Period-3": allValues[8],
            "Licensure Level of Care-3": allValues[9],
            "Beds at End of Report Period-3": allValues[10],
            "Licensed Bed Days During Report Period-3": allValues[11],
            "Beds at Beginning of Report Period-4": allValues[12],
            "Licensure Level of Care-4": allValues[13],
            "Beds at End of Report Period-4": allValues[14],
            "Licensed Bed Days During Report Period-4": allValues[15],
            "Beds at Beginning of Report Period-5": allValues[16],
            "Licensure Level of Care-5": allValues[17],
            "Beds at End of Report Period-5": allValues[18],
            "Licensed Bed Days During Report Period-5": allValues[19],
            "Beds at Beginning of Report Period-6": allValues[20],
            "Licensure Level of Care-6": allValues[21],
            "Beds at End of Report Period-6": allValues[22],
            "Licensed Bed Days During Report Period-6": allValues[23],
            "Op-Cash on Hand and in Banks":cash,
            "A.Cons-Cash on Hand and in Banks":cash_ac,
            "Op-Cash-Patient Deposits":cash_dep,
            "A.Cons-Cash-Patient Deposits":cash_dep_ac,
            "Op-Accounts & Short-Term Notes Receivable":acc,
            "A.Cons-Accounts & Short-Term Notes Receivable":acc_ac,
            "Op-Supply Inventory":supply_op,
            "A.Cons-Supply Inventory":supply_ac,
            "Op-Short-Term Investments":short_op,
            "A.Cons-Short-Term Investments":short_ac,
            "Op-Prepaid Insurance":prepaid_op,
            "A.Cons-Prepaid Insurance":prepaid_ac,
            "Op-Other Prepaid Expenses":other_op,
            "A.Cons-Other Prepaid Expenses":other_ac,
            "Op-Accounts Receivable (owners or related parties)":account_op,
            "A.Cons-Accounts Receivable (owners or related parties)":account_ac,
            "Op-Other(specify)":other_specify_op,
            "A.Cons-Other(specify)":other_specify_ac,
            "Op-TOTAL Current Assets (sum of lines 1 thru 9)":current_assets_op,
            "A.Cons-TOTAL Current Assets (sum of lines 1 thru 9)":current_assets_ac,
            "Op-Long-Term Notes Receivable":long_op,
            "A.Cons-Long-Term Notes Receivable":long_ac,
            "Op-Long-Term Investments":long_term_op,
            "A.Cons-Long-Term Investments":long_term_ac,
            "Op-Land":land_op,
            "A.Cons-Land":land_ac,
            "Op-Buildings, at Historical Cost":build_op,
            "A.Cons-Buildings, at Historical Cost":build_ac,
            "Op-Leasehold Improvements, at Historical Cost":leasehold_op,
            "A.Cons-Leasehold Improvements, at Historical Cost":leasehold_ac,
            "Op-Equipment, at Historical Cost":equipment_op,
            "A.Cons-Equipment, at Historical Cost":equipment_ac,
            "Op-Accumulated Depreciation (book methods)":accumulated_op,
            "A.Cons-Accumulated Depreciation (book methods)":accumulated_ac,
            "Op-Deferred Charges":deferred_op,
            "A.Cons-Deferred Charges":deferred_ac,
            "Op-Organization & Pre-Operating Costs":organization_op,
            "A.Cons-Organization & Pre-Operating Costs":organization_ac,
            "Op-Accumulated Amortization - Organization & Pre-Operating Costs":accumulated_amortization_op,
            "A.Cons-Accumulated Amortization - Organization & Pre-Operating Costs":accumulated_amortization_ac,
            "Op-Restricted Funds":restricted_op,
            "A.Cons-Restricted Funds":restricted_ac,
            "Op-Other Long-Term Assets (specify)":other_long_term_op,
            "A.Cons-Other Long-Term Assets (specify)":other_long_term_ac,
            "Op-Other(specify)":other_specify_two_op,
            "A.Cons-Other(specify)":other_specify_two_ac,
            "Op-TOTAL Long-Term Assets (sum of lines 11 thru 23)":long_term_assets_op,
            "A.Cons-TOTAL Long-Term Assets (sum of lines 11 thru 23)":long_term_assets_ac,
            "Op-TOTAL ASSETS (sum of lines 10 and 24)":total_assets_op,
            "A.Cons-TOTAL ASSETS (sum of lines 10 and 24)":total_assets_ac,
            "Op-Accounts Payable":accounts_payable_op,
            "A.Cons-Accounts Payable":accounts_payable_ac,
            "Op-Officer's Accounts Payable":offer_op,
            "A.Cons-Officer's Accounts Payable":offer_ac,
            "Op-Accounts Payable-Patient Deposits":accounts_patient_op,
            "A.Cons-Accounts Payable-Patient Deposits":accounts_patient_ac,
            "Op-Short-Term Notes Payable":sort_term_op,
            "A.Cons-Short-Term Notes Payable":sort_term_ac,
            "Op-Accrued Salaries Payable":accured_salaries_op,
            "A.Cons-Accrued Salaries Payable":accured_salaries_ac,
            "Op-Accrued Taxes Payable (excluding real estate taxes)":accrued_taxes_op,
            "A.Cons-Accrued Taxes Payable (excluding real estate taxes)":accrued_taxes_ac,
            "Op-Accrued Real Estate Taxes(Sch.IX-B)":accrued_real_op,
            "A.Cons-Accrued Real Estate Taxes(Sch.IX-B)":accrued_real_ac,
            "Op-Accrued Interest Payable":accrued_interest_op,
            "A.Cons-Accrued Interest Payable":accrued_interest_ac,
            "Op-Deferred Compensation":deferred_compensation_op,
            "A.Cons-Deferred Compensation":deferred_compensation_ac,
            "Op-Federal and State Income Taxes":federal_op,
            "A.Cons-Federal and State Income Taxes":federal_ac,
            "Op-TOTAL Current Liabilities (sum of lines 26 thru 37)":current_liabilities_op,
            "A.Cons-TOTAL Current Liabilities (sum of lines 26 thru 37)":current_liabilities_ac,
            "Op-Long-Term Notes Payable":long_term_notes_op,
            "A.Cons-Long-Term Notes Payable":long_term_notes_ac,
            "Op-Mortgage Payable":mortgage_op,
            "A.Cons-Mortgage Payable":mortgage_ac,
            "Op-Bonds Payable":bond_op,
            "A.Cons-Bonds Payable":bond_ac,
            "Op-Deferred Compensation":deferred_compensation_two_op,
            "A.Cons-Deferred Compensation":deferred_compensation_two_ac,
            "Op-TOTAL Long-Term Liabilities (sum of lines 39 thru 44)":long_liabilities_op,
            "A.Cons-TOTAL Long-Term Liabilities (sum of lines 39 thru 44)":long_liabilities_ac,
            "Op-TOTAL LIABILITIES (sum of lines 38 and 45)":total_liabilities_op,
            "A.Cons-TOTAL LIABILITIES (sum of lines 38 and 45)":total_liabilities_ac,
            "Op-TOTAL EQUITY(page 18, line 24)":total_equity_op,
            "A.Cons-TOTAL EQUITY(page 18, line 24)":total_equity_ac,
            "Op-TOTAL LIABILITIES AND EQUITY (sum of lines 46 and 47)":liabilities_equity_op,
            "A.Cons-TOTAL LIABILITIES AND EQUITY (sum of lines 46 and 47)":liabilities_equity_ac,
            "Gross Revenue -- All Levels of Care": gross,
            "Discounts and Allowances for all Levels":disc,
            "SUBTOTAL Inpatient Care (line 1 minus line 2)":sub_total,
            "Day Care":day_care,
            "Other Care for Outpatients":other_care,
            "Therapy": therapy,
            "Oxygen":oxygen,
            "SUBTOTAL Ancillary Revenue (lines 4 thru 7)":subtotal_ancillary,
            "Payments for Education":payment,
            "Other Government Grants":other_gov,
            "CNA Training Reimbursements":can,
            "Gift and Coffee Shop":gift,
            "Barber and Beauty Care":barber,
            "Non-Patient Meals":non_pat,
            "Telephone, Television and Radio":tele,
            "Rental of Facility Space":rental,
            "Sale of Drugs":sale_drug,
            "Sale of Supplies to Non-Patients":sale_sup,
            "Laboratory":lab,
            "Radiology and X-Ray":radio,
            "Other Medical Services":other_medi,
            "Laundry":laundry,
            "SUBTOTAL Other Operating Revenue (lines 9 thru 22":subtotal_other,
            "Contributions":contri,
            "Interest and Other Investment Income***":interest,
            "SUBTOTAL Non-Operating Revenue (lines 24 and 25)":subtotal_non_op,
            "Settlement Income (Insurance, Legal, Etc.)":settlement,
            "SUBTOTAL Other Revenue (lines 27, 28 and 28a)":subtotal_other_revenue,
            "TOTAL REVENUE (sum of lines 3, 8, 23, 26 and 29)":total_revenue,
            "General Services":general,
            "Health Care":health,
            "General Administration":general_admin,
            "Ownership":ownership,
            "Special Cost Centers":special_cost,
            "Provider Participation Fee":provider,
            "TOTAL EXPENSES (sum of lines 31 thru 39)":total_expenses,
            "Income before Income Taxes (line 30 minus line 40)":income_before,
            "Income Taxes":income_tax,
            "NET INCOME OR LOSS FOR THE YEAR (line 41 minus line 42)":net_income,
            "Medicaid - Net Inpatient Revenue":madicaid,
            "Private Pay - Net Inpatient Revenue":private,
            "Medicare - Net Inpatient Revenue":medicare,
            "Other-(specify) Med B income":other_specify_med,
            "Other-(specify)":other_speci,
            "TOTAL Inpatient Care Revenue (This total must agree to Line 3)":total_inpa
        }

        df = pd.DataFrame(dict1)
        if not os.path.exists("sheet/"):
            os.makedirs("sheet/")
       
        df.to_csv("static/data_" + str(date_name) + ".csv", index=False)

        remove_pdf_files = glob.glob('documents/*')
        for rpf in remove_pdf_files:
            os.remove(rpf)

        remove_doc_files = glob.glob('out_text/*')
        for rdf in remove_doc_files:
            os.remove(rdf)
        # if not os.path.exists("sheet/"):
        #         os.makedirs("out_text/")


        #  if form.is_valid():
        #      for f in files:
        #          file_instance = UploadPDF(upload=f)
        #          file_instance.save()
        # context = {
        #     'sheet':"data_" + str(date_name) + ".csv",
        # }
        context["sheet"] = "data_" + str(date_name) + ".csv"
    else:
         context["form"] = UploadPDFForm()
    return render(request,'demo.html',context)

